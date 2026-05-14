import os
import json
import pdfplumber
from docx import Document
from pathlib import Path


def extract_text_from_pdf(file_path):
    """Extract text from PDF using pdfplumber."""
    text = ""
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        print(f"Failed to extract from {file_path}: {e}")
        return ""
    return text


def extract_text_from_docx(file_path):
    """Extract text from DOCX files."""
    text = ""
    try:
        doc = Document(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
    except Exception as e:
        print(f"Failed to extract from {file_path}: {e}")
        return ""
    return text


def extract_text_from_txt(file_path):
    """Extract text from TXT files."""
    text = ""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
    except Exception as e:
        print(f"Failed to extract from {file_path}: {e}")
        return ""
    return text


def extract_text(file_input):
    """Auto-detect file type and extract text. Handles both file paths and Streamlit UploadedFile objects."""
    import io
    
    # Determine file extension and get file data
    if hasattr(file_input, 'name'):  # Streamlit UploadedFile object
        filename = file_input.name
        file_data = file_input.read()
        file_ext = Path(filename).suffix.lower()
    else:  # File path string
        filename = str(file_input)
        file_ext = Path(filename).suffix.lower()
        with open(file_input, 'rb') as f:
            file_data = f.read()
    
    if file_ext == '.pdf':
        return extract_text_from_pdf(io.BytesIO(file_data))
    elif file_ext == '.docx':
        return extract_text_from_docx(io.BytesIO(file_data))
    elif file_ext == '.txt':
        return file_data.decode('utf-8', errors='ignore')
    else:
        raise ValueError(f"Unsupported file format: {file_ext}")
